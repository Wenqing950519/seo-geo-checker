from __future__ import annotations

import datetime as dt
import shutil
import tempfile
import zipfile
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent / '.codex_deps'))
from docx import Document
from docx.shared import Pt
from xml.etree import ElementTree as ET
from xml.sax.saxutils import escape

ROOT = Path(__file__).resolve().parent
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
DC = "http://purl.org/dc/elements/1.1/"
DCTERMS = "http://purl.org/dc/terms/"
NS = {"w": W}
ET.register_namespace("w", W)


def qn(tag: str) -> str:
    prefix, local = tag.split(":", 1)
    return f"{{{W if prefix == 'w' else prefix}}}{local}"


def paragraph(text: str = "", style: str | None = None, bold: bool = False):
    p = ET.Element(qn("w:p"))
    if style:
        ppr = ET.SubElement(p, qn("w:pPr"))
        ET.SubElement(ppr, qn("w:pStyle"), {qn("w:val"): style})
    if text:
        r = ET.SubElement(p, qn("w:r"))
        if bold:
            rpr = ET.SubElement(r, qn("w:rPr"))
            ET.SubElement(rpr, qn("w:b"))
        t = ET.SubElement(r, qn("w:t"))
        t.text = text
    return p


def page_break():
    p = ET.Element(qn("w:p"))
    r = ET.SubElement(p, qn("w:r"))
    ET.SubElement(r, qn("w:br"), {qn("w:type"): "page"})
    return p


def bullet(text: str):
    return paragraph(f"• {text}")


def append_addendum(source: Path, target: Path, title: str, sections: list[tuple[str, list[str]]]):
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document(str(source))
    if "Technical_Whitepaper" in target.name:
        replacements = {
            "V2.1": "V3.0",
            "2.0 / Current Implementation Whitepaper": "3.0 / Perplexity-first GEO Whitepaper",
            "2026-07-14": "2026-07-16",
        }
        paragraphs = list(document.paragraphs)
        for section in document.sections:
            paragraphs.extend(section.header.paragraphs)
            paragraphs.extend(section.footer.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        for paragraph in paragraphs:
            for run in paragraph.runs:
                for old, new in replacements.items():
                    if old in run.text:
                        run.text = run.text.replace(old, new)
    document.add_page_break()

    def add_compact(text: str, style: str | None = None, size: float = 9.5):
        paragraph = document.add_paragraph(style=style)
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        return paragraph

    title_paragraph = add_compact(title, "Title", 18)
    title_paragraph.paragraph_format.space_after = Pt(5)
    add_compact("更新日期：2026-07-16", None, 9)
    for heading, lines in sections:
        heading_paragraph = add_compact(heading, "Heading 1", 13)
        heading_paragraph.paragraph_format.space_before = Pt(5)
        heading_paragraph.paragraph_format.space_after = Pt(2)
        for line in lines:
            if line.startswith("## "):
                add_compact(line[3:], "Heading 2", 11)
            else:
                add_compact(line, None, 9.5)
    document.core_properties.title = title
    document.core_properties.modified = dt.datetime.now(dt.timezone.utc)
    document.save(str(target))

COMMON = [
    ("V3 量測架構", [
        "GEO 主分數由 Perplexity 搜尋觀測 50%、內容可引用性 30%、必要技術存取 20% 組成。",
        "每個網站固定使用 3 次 Perplexity：1 次精確實體驗證與 2 次非品牌探索。",
        "Gemini Flash-Lite 每站使用 1 次，只標準化基本資訊、產業、網站結構與內容特徵；不參與計分。",
        "Perplexity 證據不足時 GEO 分數為未知，不以站內準備度或 Gemini 判讀補分。",
    ]),
    ("白皮書研究模式", [
        "研究 Skill 與正式網站共同匯入 mock-api/lib/geo-measurement.js，避免權重與查詢邏輯漂移。",
        "批次執行要求 Perplexity 與 Gemini 兩個獨立硬上限；超過任一上限時，第一個付費呼叫前即停止。",
        "輸出包含 JSONL、CSV、統計摘要、方法檔、查詢觀測、來源網址、模型、失敗狀態與資料集 SHA-256。",
        "Gemini 研究 schema 不含建議、行動、改寫、改善或預期成效欄位。",
    ]),
    ("最終驗證與限制", [
        "2026-07-16 實測：壽司郎 GEO 65、站內準備度 63；Hunterest GEO 41、站內準備度 87。",
        "結果證明站內結構完整不再自動換成高 GEO 分；外部實體、品牌提及與官網引用證據才是主差異。",
        "量測只代表指定 Perplexity 模型、固定查詢集與收集期間，不代表所有 AI 引擎的普遍排名。",
        "最終資料集 SHA-256：83d50e17a6962342d5e1baf7c44ea3e97fb40bda0c4ff2ba7703b8304010cd6c。",
    ]),
]

TECH = COMMON + [
    ("執行與安全", [
        "本機 Gemini 若因出口地區遭拒，研究腳本可改走 Render 上的受保護代理 POST /api/internal/research-profile。",
        "代理必須帶 X-Admin-Token，未授權請求回傳 401；API key 不傳回前端或資料集。",
        "npm.cmd test 會驗證演算法邊界、網站與 Skill 綁定、禁止建議欄位與代理密碼契約。",
    ]),
]

BRD = COMMON + [
    ("商業影響", [
        "白皮書資料從零 API 站內檢查升級為具 Perplexity 搜尋來源的實證資料，可提高渠道合作與研究引用信任度。",
        "每站固定 3+1 次模型呼叫，使 300 或 400 站研究可在執行前精確估算請求數。",
        "不產生逐站優化建議，避免白皮書批次成本被長篇輸出放大，也降低未經同意公開小商家弱點的風險。",
    ]),
]

MRD = COMMON + [
    ("市場研究口徑", [
        "白皮書可報告品牌提及率、官網引用率、實體對齊率與 GEO 分布。",
        "不得將 Perplexity 結果外推為 ChatGPT、Gemini、Claude 或所有 AI 搜尋引擎的共同結果。",
        "每份發布內容需揭露樣本數、模型、查詢集、收集期間、失敗率與資料集雜湊。",
    ]),
]

PRD = COMMON + [
    ("產品需求與驗收條件", [
        "• 批次腳本需接受 TXT/CSV、去重網址、JSONL 續跑、雙模型硬上限與 1 至 4 的並行數。",
        "• 每筆結果需包含 geo_score、site_readiness_score、perplexity_score、mention_rate、official_citation_rate、source_urls、Gemini profile 與 evidence_hash。",
        "• Gemini profile 不得包含 recommendation、action、rewrite、improvement 或 expected impact。",
        "• 網站與 Skill 不得各自保存一份權重；同步測試失敗時禁止部署。",
        "• 研究代理必須驗證 X-Admin-Token，無密碼請求回傳 401。",
    ]),
]

PORTFOLIO = COMMON + [
    ("作品集更新重點", [
        "產品由技術 SEO 打分器轉為 Perplexity-first GEO 量測平台。",
        "差異化來自可追溯查詢、來源、實體對齊與固定成本研究流程，而不是單純的大模型文案生成。",
        "共享核心與測試契約讓 SaaS 報告、研究 Skill、白皮書資料保持同一演算法版本。",
    ]),
]


def main():
    jobs = [
        (ROOT / "GEOCheck_Technical_Whitepaper_V2.1_2026-07-14.docx", ROOT / "GEOCheck_Technical_Whitepaper_V3_2026-07-16.docx", "GeoCheck 技術白皮書 V3.0", TECH),
        (ROOT / "business_docs/GEOCheck_BRD_2026-07-14.docx", ROOT / "business_docs/GEOCheck_BRD_2026-07-16.docx", "GeoCheck BRD V3 更新版", BRD),
        (ROOT / "business_docs/GEOCheck_MRD_2026-07-14.docx", ROOT / "business_docs/GEOCheck_MRD_2026-07-16.docx", "GeoCheck MRD V3 更新版", MRD),
        (ROOT / "business_docs/GEOCheck_PRD_2026-07-14.docx", ROOT / "business_docs/GEOCheck_PRD_2026-07-16.docx", "GeoCheck PRD V3 更新版", PRD),
        (ROOT / "GEO_Intelligence_Platform_PRD_Portfolio.docx", ROOT / "GEO_Intelligence_Platform_PRD_Portfolio_V3_2026-07-16.docx", "GEO Intelligence Platform 作品集 PRD V3", PORTFOLIO),
    ]
    for source, target, title, sections in jobs:
        append_addendum(source, target, title, sections)
        with zipfile.ZipFile(target, "r") as check:
            bad = check.testzip()
            if bad:
                raise RuntimeError(f"Corrupt DOCX entry in {target}: {bad}")
        print(target.relative_to(ROOT))


if __name__ == "__main__":
    main()
