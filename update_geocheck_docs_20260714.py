from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
NAVY = "0B2545"
BLUE = "2E74B5"
GRAY = "5E6B78"
PALE = "E8EEF5"


def set_font(run, size=10, bold=False, color=None):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def replace_in_container(container):
    replacements = [
        ("Agnes 2.0 Flash", "Gemini 3.1 Flash-Lite"),
        ("Agnes", "Gemini 3.1 Flash-Lite"),
        ("Brave Search", "Perplexity Sonar"),
        ("Brave", "Perplexity"),
        ("V2.0", "V2.1"),
    ]
    for paragraph in container.paragraphs:
        for run in paragraph.runs:
            for old, new in replacements:
                run.text = run.text.replace(old, new)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_container(cell)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_font(run, size=18, bold=True, color=NAVY)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=10.2, color="1F2933")
    return p


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["項目", "目前狀態", "判定依據"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cell.paragraphs[0].add_run(text), size=9.2, bold=True, color="FFFFFF")
    for item, state, evidence in rows:
        cells = table.add_row().cells
        for index, text in enumerate((item, state, evidence)):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_font(cells[index].paragraphs[0].add_run(text), size=8.8, color="1F2933")
    for row in table.rows:
        row.cells[0].width = Inches(1.55)
        row.cells[1].width = Inches(1.35)
        row.cells[2].width = Inches(3.6)
    return table


def add_update_page(doc, document_type):
    doc.add_page_break()
    add_title(doc, f"2026-07-14 實作狀態更新（{document_type}）")
    add_body(doc, "本頁覆蓋本文件中較早的供應商名稱、版本敘述與待驗證假設；以 2026-07-14 的程式碼、測試與供應商連線檢查為準。任何未公開或未驗證的 AI 引用行為一律標示為未知。")
    add_status_table(doc, [
        ("規則評分與爬蟲", "已驗證", "核心分數由多頁抓取、技術訊號與確定性規則決定；自動測試通過。"),
        ("Perplexity Sonar", "已啟用", "真實供應商連線成功；用於網站公開脈絡與來源探索，不改寫核心分數。"),
        ("Gemini 3.1 Flash-Lite", "部署待驗證", "程式與設定已接通，但目前執行位置收到 HTTP 400：User location is not supported for the API use。"),
        ("模型失敗降級", "已驗證", "Gemini 不可用時回傳本地確定性報告，AI 定位標記為未驗證。"),
        ("成本與私有後台", "已驗證", "已記錄 token、延遲、狀態與預估成本；私有路徑加密碼雙層保護。"),
    ])
    if document_type == "BRD":
        add_body(doc, "商業決策：在 Gemini 尚未於正式部署區域驗證前，對外銷售應定位為「AI 搜尋可讀性／技術整備度健檢與修正」，不可承諾 Gemini 引用、曝光或流量成果。成本台帳已可支援每單 API 成本與毛利的實際記錄。")
    elif document_type == "MRD":
        add_body(doc, "市場訊息：小商家需要可理解、可執行的網站修正，而非無法驗證的排名承諾。產品競爭點是繁體在地化診斷、可追溯證據與低成本修正流程；跨引擎 AI 引用監測仍屬後續驗證能力。")
    elif document_type == "PRD":
        add_body(doc, "PRD 驗收條件新增：正式部署後，`POST /api/test-provider` 必須成功回傳 Gemini 模型與延遲；若未成功，系統必須維持本地降級與明確限制說明。私有成本後台使用 `ADMIN_PATH_TOKEN` 與 `ADMIN_TOKEN`，不提供公開入口。")
    else:
        add_body(doc, "供應商分工：Perplexity Sonar 已驗證為公開脈絡補充；Gemini 3.1 Flash-Lite 僅在部署區域通過連線後才執行語意解讀。兩者都不得改寫核心分數。成本台帳為 JSONL 持久化快照；非持久化主機需移至資料庫或持久化磁碟。")
    add_body(doc, "上線前條件：部署平台必須設定 GEMINI_API_KEY、PERPLEXITY_API_KEY、ADMIN_PATH_TOKEN、ADMIN_TOKEN；Gemini 請求需位於支援區域或改用 Vertex AI。成本單價為預估參數，須以供應商帳單校正。")


def update_one(source, target, document_type):
    doc = Document(source)
    replace_in_container(doc)
    for section in doc.sections:
        replace_in_container(section.header)
        replace_in_container(section.footer)
    add_update_page(doc, document_type)
    doc.core_properties.subject = "GeoCheck V2.1 current implementation and deployment status"
    doc.core_properties.comments = "Updated 2026-07-14 with verified provider, deployment, cost, and private admin status."
    doc.save(target)
    print(target)


def main():
    update_one(ROOT / "GEOCheck_Technical_Whitepaper_V2_2026-07-13.docx", ROOT / "GEOCheck_Technical_Whitepaper_V2.1_2026-07-14.docx", "技術白皮書")
    business = ROOT / "business_docs"
    update_one(business / "GEOCheck_BRD_2026-07-13.docx", business / "GEOCheck_BRD_2026-07-14.docx", "BRD")
    update_one(business / "GEOCheck_MRD_2026-07-13.docx", business / "GEOCheck_MRD_2026-07-14.docx", "MRD")
    update_one(business / "GEOCheck_PRD_2026-07-13.docx", business / "GEOCheck_PRD_2026-07-14.docx", "PRD")


if __name__ == "__main__":
    main()
